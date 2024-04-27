import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

def format_class_year(class_year):
    if len(class_year) == 4:
        return class_year[-2:]
    elif len(class_year) == 2:
        return class_year
    else:
        return "XX"

def generate_acknowledgment_paragraph(cadet_info):
    paragraphs = []
    for info in cadet_info:
        name, company, class_year, assistance_type, problem_description, modification_description, place_and_date = info
        formatted_name = f"{name} CDT {company[:1]}-{company[1:]} '{format_class_year(class_year)}"
        if place_and_date:
            place_and_date = f"West Point, NY, {place_and_date}"
        if problem_description and not problem_description.endswith('.'):
            problem_description += '.'
        if modification_description and not modification_description.endswith('.'):
            modification_description += '.'
        paragraph = f"{formatted_name}. Assistance given to the author, {assistance_type}. {problem_description} {modification_description} {place_and_date}"
        paragraphs.append(paragraph)
    return paragraphs

def generate_ai_acknowledgment_paragraph(ai_info):
    ai_paragraphs = []
    for info in ai_info:
        ai_type, other_ai_type, prompt, ai_usage, ai_link, date = info
        if ai_type == "Other":
            ai_type = other_ai_type
        if not ai_usage.endswith('.'):
            ai_usage += '.'
        ai_link = f"({ai_link})"
        paragraph = f"{ai_type}. Assistance given to the author, AI. I used the following prompt in {ai_type}: \"{prompt}\". {ai_usage} {ai_link}. West Point, NY, {date}."
        ai_paragraphs.append(paragraph)
    return ai_paragraphs

def create_word_document(paragraphs):
    doc = Document()

    # Set margins to 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set line spacing for entire document
    doc.styles['Normal'].paragraph_format.line_spacing = 1.0

    # Add title with space before first paragraph
    title = doc.add_paragraph("Acknowledgment of Assistance")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.size = Pt(12)
    doc.add_paragraph()  # Add space between title and first paragraph

    # Add acknowledgment paragraphs with space in between
    for paragraph in sorted(paragraphs):
        p = doc.add_paragraph(paragraph)
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(12)
        doc.add_paragraph()  # Add space between paragraphs

    # Save the document to a BytesIO object
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes

def main():
    st.set_page_config(page_title="APG", page_icon="📄")

    st.title('Acknowledgment Page Generator')

    tabs = ["Assistant Information", "Artificial Intelligence"]
    selected_tab = st.sidebar.radio("Assistance Source", tabs)

    if selected_tab == "Assistant Information":
        st.sidebar.header("Assistant Information")
        name = st.sidebar.text_input("Name (Last, First):")
        company = st.sidebar.text_input("Company: (i.e. C4)")
        class_year = st.sidebar.text_input("Class Year:")
        assistance_type = st.sidebar.selectbox("Assistance Type:", ["Verbal", "Written", "E-mail", "Other"])
        problem_description = st.sidebar.text_area("The exact portion/problem(s) which assistance was received:")
        modification_description = st.sidebar.text_area("How you used that assistance to modify your work:")
        place_and_date = st.sidebar.text_input("Date (DDMMMYYYY):")

        if problem_description and not problem_description.endswith('.'):
            problem_description += '.'
        if modification_description and not modification_description.endswith('.'):
            modification_description += '.'

        cadet_info = [
            (name, company, class_year, assistance_type, problem_description, modification_description, place_and_date)
        ]

        if st.sidebar.button("Add Acknowledgment"):
            st.session_state["acknowledgments"] = st.session_state.get("acknowledgments", [])
            st.session_state["acknowledgments"].extend(generate_acknowledgment_paragraph(cadet_info))

    elif selected_tab == "Artificial Intelligence":
        st.sidebar.header("Artificial Intelligence Information")
        ai_type = st.sidebar.selectbox("What AI did you use?", ["ChatGPT", "Other"])
        if ai_type == "Other":
            other_ai_type = st.sidebar.text_input("Please specify the type of AI used:")
        else:
            other_ai_type = ""
        prompt = st.sidebar.text_area("What prompt did you use?")
        ai_usage = st.sidebar.text_area("How did you use the AI's assistance?")
        ai_link = st.sidebar.text_input("Link to the AI assistance chat:")
        date = st.sidebar.text_input("Date (DDMMMYYYY):")

        ai_info = [
            (ai_type, other_ai_type, prompt, ai_usage, ai_link, date)
        ]

        if st.sidebar.button("Add AI Acknowledgment"):
            st.session_state["acknowledgments"] = st.session_state.get("acknowledgments", [])
            st.session_state["acknowledgments"].extend(generate_ai_acknowledgment_paragraph(ai_info))

    if "acknowledgments" in st.session_state:
        acknowledgment_paragraphs = st.session_state["acknowledgments"]
        if len(acknowledgment_paragraphs) > 0:
            if st.button("Delete Last Acknowledgment"):
                del acknowledgment_paragraphs[-1]

        # Display acknowledgment paragraphs
        for paragraph in acknowledgment_paragraphs:
            st.markdown(paragraph)

        # Create a Word document with the acknowledgment paragraphs
        word_doc = create_word_document(acknowledgment_paragraphs)

        # Provide a download button for the Word document
        st.download_button(
            label="Download Acknowledgment Paragraph",
            data=word_doc,
            file_name="acknowledgment_paragraph.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Add footer
    st.markdown("""
    <div style="position: fixed; bottom: 10px; right: 10px; font-size: 14px; color: gray;">
        Created by CDT Kim, Juhun '27.
    </div>
    """, unsafe_allow_html=True)

if __name__ == '__main__':
    main()
